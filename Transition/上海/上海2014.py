# -*- coding: utf-8 -*-
"""
Created on Thu Apr  7 22:19:17 2022
@author: TongJi|NWPU wjm
No pains,no gains. 
You must to have the ability to protect family mumbers.
"""
from docx import Document
import re
import pandas as pd
import xlwt
import xlrd
import os

def date(stamp):
    """输出时间格式
    """
    delta = pd.Timedelta(str(stamp)+'D')
    real_time = pd.to_datetime('1899-12-30') + delta
    return real_time

def word_extraction_shanghai_table(province:str,year:str):
    
    path = r'E:\Spyderworkplace\excel\总实验数据'
    path_output = r'E:\Spyderworkplace\excel\修正数据'
    filelist = os.listdir(path)
    """ 生成excel的表头
    """
    i_out = 1
    workbook = xlwt.Workbook(encoding='utf-8')
    worksheet = workbook.add_sheet('sheet1', cell_overwrite_ok=True)
    worksheet.write(0, 0, label='年度')
    worksheet.write(0, 1, label='通知发布日期')
    worksheet.write(0, 2, label='来源公告名称')
    worksheet.write(0, 3, label='事务所所在省份')
    worksheet.write(0, 4, label='事务所全称')
    worksheet.write(0, 5, label='总所or分所')
    worksheet.write(0, 6, label='总/分所所长（如有）')
    worksheet.write(0, 7, label='注册会计师姓名')
    worksheet.write(0, 8, label='注册会计师ID')
    worksheet.write(0, 9, label='备注')
    for file in filelist:
        if '.docx' in file and year in file and province in file:  # 搜索docx
            word_path = os.path.join(path, file)
            shiwusuo_number = getSumnumber(word_path)
            shiwusuo_title = getTitle(word_path)
            shiwusuo = getTableValue(word_path)
            newshiwusuo = chongzu(shiwusuo,shiwusuo_number)
            wb_read = xlrd.open_workbook(os.path.join(path,'源地址.xls'))  # 从源地址excel提取'通知发布日期' '来源公告名称'
            for i in range(wb_read.sheet_by_index(0).nrows - 1):
                if year in wb_read.sheet_by_index(0).cell_value(i, 0) and province in wb_read.sheet_by_index(
                        0).cell_value(i, 0):
                    d = wb_read.sheet_by_index(0).cell(i, 1).value
                    date_pb = str(date(d)).split(" ")[0]
                    name_pb = wb_read.sheet_by_index(0).cell_value(i, 0)


            for st,x in enumerate(shiwusuo_title):
                print(st)
                for y in newshiwusuo[st]:
                    worksheet.write(i_out, 0, year)
                    worksheet.write(i_out, 1, date_pb)
                    worksheet.write(i_out, 2, name_pb)
                    worksheet.write(i_out, 3, province)
                    worksheet.write(i_out, 4, x)  # 事务所
                    if '总所' in x:
                        worksheet.write(i_out, 5, '总所')
                    elif '分所' in x or '分公司' in x:
                        worksheet.write(i_out, 5, '分所')
                    else:
                        worksheet.write(i_out, 5, '总所')
                    worksheet.write(i_out, 7, y)
                    i_out = i_out + 1
                    
    workbook.save(os.path.join(path_output,province+year+'.xls'))   
           


def checkErrorTable(path): #检测跨行连接的情况
    doc = Document(path)
    for i in doc.paragraphs:

        if (re.search(r'事务',i.text) or re.search(r'分所',i.text) or (re.search(r'代管', i.text)))  and '市' not in i.text and '以下' not in i.text:
            continue
        elif (i.text != " " and i.text != "" and (not re.search(r'[0-9]+',i.text))):
            print(i.text)

def getTitle(path:str):#检测事务所名称
    doc = Document(path)
    title = []
    for i in doc.paragraphs:
        j = i.text
        if (re.search(r'事务',j) or re.search(r'分所',j) or (re.search(r'代管', j)))  and '市' not in j and '以下' not in j:
        #print(paragraph.split()[0])
            j = j.split()[0]
            j = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s]+', "",j)  # 事务所名称
            j = re.sub('人', '', j)
            j = re.sub('（', '', j)
            j = re.sub('）', '', j)
            j = j.split()[0]
            title.append(j)
    return title

def getTableValue(path:str):#读取各事务所中会计师姓名
    single_word = []
    doc = Document(path)
    table = list(doc.tables)
    shiwusuo_name_tables = []
    for u,t in enumerate(table):
        row = list(t.rows)
        shiwusuo_name_rows = []
        for v,r in enumerate(row):
            cell = list(r.cells)
            i = 0
            while i < len(cell):
                if(i < len(cell)-1):
                    if(cell[i].text.replace(" ","").replace('\n','').replace('\t','') == cell[i+1].text.replace(" ","").replace('\n','').replace('\t','')):
                        shiwusuo_name_rows.append(cell[i].text.replace(" ","").replace('\n',''))
                        for j in range(i+1,len(cell)):
                            k = i+2
                            if(cell[i].text.replace(" ","").replace('\n','').replace('\t','') != cell[j].text.replace(" ","").replace('\n','').replace('\t','')):
                                k = j #找到不相同的索引值
                                break
                            else:
                                continue
                        i = k
                        continue
                    elif(len(cell[i].text.replace(" ","").replace('\n','').replace('\t','')) == len(cell[i+1].text.replace(" ","").replace('\n','').replace('\t','')) and len(cell[i].text.replace(" ","").replace('\n','').replace('\t','')) == 1):
                        shiwusuo_name_rows.append(cell[i].text.replace(" ","").replace('\n','').replace('\t','') + cell[i+1].text.replace(" ","").replace('\n','').replace('\t',''))
                        i = i + 2
                        continue
                    else:
                        shiwusuo_name_rows.append(cell[i].text.replace(" ","").replace('\n','').replace('\t',''))
                else:
                    shiwusuo_name_rows.append(cell[i].text.replace(" ","").replace('\n','').replace('\t',''))
                i = i + 1
                
        shiwusuo_name_rows = splitRepeat(shiwusuo_name_rows)
        
        for i,name in enumerate(shiwusuo_name_rows):
            if len(name) == 1 or len(name) > 3 :
                single_word.append(str(u)+'_'+str(i)+'_'+name)
                
        shiwusuo_name_tables.append(shiwusuo_name_rows)
        
    return shiwusuo_name_tables

def splitRepeat(single_tab : list):#去除事务所中异常的会计师姓名
    label = False
        
    for i in range(len(single_tab)-8):
        if single_tab[i] == single_tab[i+8]:
            single_tab = single_tab[0:i+8]
            break
        
    for i in range(len(single_tab)-1):
        if single_tab[i] == single_tab[i+1]:
            single_tab = single_tab[0:i+1]
            break
        
    for i in range(len(single_tab)):
        if single_tab[i] == '':
            single_tab[i] = '空值'
            label = True
            
    if label:
        single_tab.remove("空值")
    return single_tab

def getSumnumber(path:str):#事务所会计师总数
    doc = Document(path)
    shiwusuo_number = []
    for i in doc.paragraphs:
        j = i.text
        if (re.search(r'事务',j) or re.search(r'分所',j) or (re.search(r'代管', j)))  and '市' not in j and '以下' not in j:
        #print(paragraph.split()[0])
            j = j.split()[0]
            j = int(re.search('[0-9]+',j).group(0)) #返回人数
            shiwusuo_number.append(j)
    return shiwusuo_number
        

def chongzu(old_shiwusuo_table : list ,number_list:list): #合并跨页的事务所会计师
    new_shiwusuo_table = []
    temp1 = old_shiwusuo_table
    for i in range(len(number_list)):
        sum_number = 0
        for j in range(len(temp1)):
            temp2 = []
            sum_number = len(temp1[j]) + sum_number
            if sum_number == number_list[i]:
                for u in range(j+1):
                    temp2 = temp2 + temp1[u]
                new_shiwusuo_table.append(temp2)
                temp1 = temp1[j+1 : len(temp1)]
                break
             
    return new_shiwusuo_table
             
                     
if __name__ == "__main__":
    '''
    path = '上海市2014年度通过任职资格检查注册会计师名单.docx'
    #checkErrorTable(path)
    title = getTitle(path)
    table,single = getTableValue(path)
    number = getSumnumber(path)
    newshiwsuo = chongzu(table, number)
    '''
    word_extraction_shanghai_table('上海', '2014')

