import os
from docx import Document
import xlwt
import xlrd
import re
import pandas as pd

def tableToList(path):
    doc = Document(path)
    tables = []
    for i,table in enumerate(doc.tables):
        rows = []
        for r,row in enumerate(table.rows):
            cells = []
            for c,cell in enumerate(row.cells):
                cells.append(cell.text)
            rows.append(cells)
        tables.append(rows)
    return tables
        
def paragraphToList(path):
    doc = Document(path)
    paragraph = doc.paragraphs
    paragraph2text = []
    for i in paragraph: #除去空行
        if i.text == '':
            continue
        else:
            paragraph2text.append(i.text)
            
        
    return paragraph2text

def getTitle(paragraph:list):
    title = []
    index = []
    for p,i in enumerate(paragraph): #除去空行
        if (re.search(r'[0-9]、', i[0:6]) and i[0:4] != '办公地址') or '、其他' in i:
            title.append(i)
            index.append(p)
        else:
            continue
    return title,index

def getShiwusuo(paragraph:list,index:list):
    shiwusuo = []
    shiwusuo = []
    i = 0
    while i < len(index):
        if i == len(index)-1:
            shiwusuo.append(paragraph[index[i]:len(paragraph)])
        else:
            shiwusuo.append(paragraph[index[i]:index[i+1]])
        i = i + 1
        
    return shiwusuo

def getPart1(shiwusuo:list):
    final_shiwusuo = []
    for i,sws in enumerate(shiwusuo):
        shiwusuo_name = [] #事务所名称
        shiwusuo_number = [] #事务所总人数
        shiwusuo_fensuo =[] #事务所是总所还是分所
        shiwusuo_manager =[] #事务所负责人:如有
        shiwusuo_single = []
        if len(sws) == 2:
            name = ""
            temp_sws = sws[0].split('\n')
            #获得事务所名称
            name = temp_sws[0]
            
            name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
            name = re.sub('人', '',name)
            name = re.sub('（）', '',name)
            shiwusuo_name.append(name)
            
            #获得事务所人数
            shiwusuo_number.append(re.search(r'[0-9]+',sws[1]).group(0))
            
            #获得事务所分类
            if '分所' in temp_sws[0] or '分公司' in temp_sws[0]:
                shiwusuo_fensuo.append('分所')
            else: 
                shiwusuo_fensuo.append('总所')
            
                
                
            #获得负责人姓名（如有）
            if '主任会计师' in sws[0] :
                index = sws[0].index('主任会计师') 
                shiwusuo_manager.append(sws[0][index+7:index+10].replace(' ',''))
            elif '合伙人' in sws[0] :
                index = sws[0].index('合伙人')
                shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
            elif '负责人' in sws[0]:
                index = sws[0].index('负责人')
                shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
            else:
                shiwusuo_manager.append('缺失')
            
            shiwusuo_single.append(shiwusuo_name)
            shiwusuo_single.append(shiwusuo_number)
            shiwusuo_single.append(shiwusuo_fensuo)
            shiwusuo_single.append(shiwusuo_manager)
            
            if '执业证书' not in sws[0]: #检查非姓名信息是否存于table中
                shiwusuo_single.append('缺失')
                
            final_shiwusuo.append(shiwusuo_single)
            
        elif len(sws) == 3:
            boolean1 = False
            #测试第三行是不是城市信息
            liststr = ['一、','二、','三、','四、','五、','六、','七、','八、','九、','十、']
            for i in liststr:
                if re.search(i,sws[2]):
                   boolean1 = True
                   
            if boolean1: #如果第三行是城市信息
                name = ""
                temp_sws = sws[0].split('\n')
                #获得事务所名称
                name = temp_sws[0]
                
                name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s]+', "",name)  # 事务所名称
                name = re.sub('人', '',name)
                name = re.sub('（）', '',name)
                
                shiwusuo_name.append(name)
                #获得事务所人数
                shiwusuo_number.append(re.search(r'[0-9]+',sws[1]).group(0))
                
                #获得事务所分类
                if '分所' in temp_sws[0] or '分公司' in temp_sws[0]:
                    shiwusuo_fensuo.append('分所')
                else:
                    shiwusuo_fensuo.append('总所')
                    
                #获得负责人姓名（如有）
                if '主任会计师' in sws[0] :
                    index = sws[0].index('主任会计师') 
                    shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
                elif '合伙人' in sws[0] :
                    index = sws[0].index('合伙人')
                    shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
                elif '负责人' in sws[0]:
                    index = sws[0].index('负责人')
                    shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
                else:
                    shiwusuo_manager.append('缺失')
                
                shiwusuo_single.append(shiwusuo_name)
                shiwusuo_single.append(shiwusuo_number)
                shiwusuo_single.append(shiwusuo_fensuo)
                shiwusuo_single.append(shiwusuo_manager)
                if '执业证书' not in sws[0]: #检查非姓名信息是否存于table中
                    shiwusuo_single.append('缺失')
                final_shiwusuo.append(shiwusuo_single)
            else: #如果第三行不是城市信息
                if re.search('执业证书',sws[1]): #检查非事务所内容的位置
                    name = ""
                    #获得事务所名称
                    name = sws[0]
                    name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
                    name = re.sub('人', '',name)
                    name = re.sub('（）', '',name)
                    
                    shiwusuo_name.append(name)
                    #获得事务所人数
                    shiwusuo_number.append(re.search(r'[0-9]+',sws[2]).group(0))
                    
                    #获得事务所分类
                    if '分所' in sws[0] or '分公司' in sws[0]:
                        shiwusuo_fensuo.append('分所')
                    else:
                        shiwusuo_fensuo.append('总所')
                        
                    #获得负责人姓名（如有）
                    if '主任会计师' in sws[1] :
                        index = sws[1].index('主任会计师') 
                        shiwusuo_manager.append(sws[1][index+7:index+10].replace(' ',''))
                    elif '合伙人' in sws[1] :
                        index = sws[1].index('合伙人')
                        shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
                    elif '负责人' in sws[1]:
                        index = sws[1].index('负责人')
                        shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
                    else:
                        shiwusuo_manager.append('缺失')
                    
                    shiwusuo_single.append(shiwusuo_name)
                    shiwusuo_single.append(shiwusuo_number)
                    shiwusuo_single.append(shiwusuo_fensuo)
                    shiwusuo_single.append(shiwusuo_manager)
                    final_shiwusuo.append(shiwusuo_single)    
                elif re.search('执业证书',sws[0]):
                    name = ""
                    temp_sws = sws[0].split('\n')
                    #获得事务所名称
                    name = temp_sws[0]
                    
                    name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
                    name = re.sub('人', '',name)
                    name = re.sub('（）', '',name)
                    
                    shiwusuo_name.append(name)
                    #获得事务所人数
                    shiwusuo_number.append(re.search(r'[0-9]+',sws[2]).group(0))
                    
                    #获得事务所分类
                    if '分所' in temp_sws[0] or '分公司' in temp_sws[0]:
                        shiwusuo_fensuo.append('分所')
                    else:
                        shiwusuo_fensuo.append('总所')
                        
                    #获得负责人姓名（如有）
                    if '主任会计师' in sws[0] :
                        index = sws[0].index('主任会计师') 
                        shiwusuo_manager.append(sws[0][index+7:index+10].replace(' ',''))
                    elif '合伙人' in sws[0] :
                        index = sws[0].index('合伙人')
                        shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
                    elif '负责人' in sws[0]:
                        index = sws[0].index('负责人')
                        shiwusuo_manager.append(sws[0][index+5:index+8].replace(' ',''))
                    else:
                        shiwusuo_manager.append('缺失')
                    
                    shiwusuo_single.append(shiwusuo_name)
                    shiwusuo_single.append(shiwusuo_number)
                    shiwusuo_single.append(shiwusuo_fensuo)
                    shiwusuo_single.append(shiwusuo_manager)
                    final_shiwusuo.append(shiwusuo_single)
                else: #事务所内容在Table中
                    name = ""
                    #获得事务所名称
                    name = sws[0]
                    name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
                    name = re.sub('人', '',name)
                    name = re.sub('（）', '',name)
                    shiwusuo_name.append(name)
                    shiwusuo_number.append(re.search(r'[0-9]+',sws[2]).group(0))
                    
                    if '分所' in sws[0] or '分公司' in sws[0]:
                        shiwusuo_fensuo.append('分所')
                    else:
                        shiwusuo_fensuo.append('总所')
                        
                    shiwusuo_manager.append('缺失')
                    
                    shiwusuo_single.append(shiwusuo_name)
                    shiwusuo_single.append(shiwusuo_number)
                    shiwusuo_single.append(shiwusuo_fensuo)
                    shiwusuo_single.append(shiwusuo_manager)
                    shiwusuo_single.append('缺失')
                    final_shiwusuo.append(shiwusuo_single)
                    
        elif len(sws) == 4:
            name = ""
            #获得事务所名称
            name = sws[0]
            name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
            name = re.sub('人', '',name)
            name = re.sub('（）', '',name)
            shiwusuo_name.append(name)
            shiwusuo_number.append(re.search(r'[0-9]+',sws[3]).group(0))
            
            if '分所' in sws[0] or '分公司' in sws[0]:
                shiwusuo_fensuo.append('分所')
            else:
                shiwusuo_fensuo.append('总所')
                
            shiwusuo_manager.append('缺失')
            
            shiwusuo_single.append(shiwusuo_name)
            shiwusuo_single.append(shiwusuo_number)
            shiwusuo_single.append(shiwusuo_fensuo)
            shiwusuo_single.append(shiwusuo_manager)
            if '执业证书' not in sws[1]: #检查非姓名信息是否存于table中
                shiwusuo_single.append('缺失')
            final_shiwusuo.append(shiwusuo_single)
            
        elif len(sws) == 5:
            name = ""
            temp_sws = sws[0].split('\n')
            #获得事务所名称
            name = temp_sws[0]
            
            name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
            name = re.sub('人', '',name)
            name = re.sub('（）', '',name)
            shiwusuo_name.append(name)
            
            #获得事务所人数
            shiwusuo_number.append(re.search(r'[0-9]+',sws[4]).group(0))
            
            #获得事务所分类
            if '分所' in temp_sws[0] or '分公司' in temp_sws[0]:
                shiwusuo_fensuo.append('分所')
            else:
                shiwusuo_fensuo.append('总所')
            
            #获得负责人姓名（如有）
            if '主任会计师' in sws[1] :
                index = sws[1].index('主任会计师') 
                shiwusuo_manager.append(sws[1][index+7:index+10].replace(' ',''))
            elif '合伙人' in sws[1] :
                index = sws[1].index('合伙人')
                shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
            elif '负责人' in sws[1]:
                index = sws[1].index('负责人')
                shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
            else:
                shiwusuo_manager.append('缺失')
            
            shiwusuo_single.append(shiwusuo_name)
            shiwusuo_single.append(shiwusuo_number)
            shiwusuo_single.append(shiwusuo_fensuo)
            shiwusuo_single.append(shiwusuo_manager)
            if '执业证书' not in sws[1]: #检查非姓名信息是否存于table中
                shiwusuo_single.append('缺失')
            final_shiwusuo.append(shiwusuo_single)
     
        elif len(sws) == 6:
            name = ""
            temp_sws = sws[0].split('\n')
            #获得事务所名称
            name = temp_sws[0]
            
            name = re.sub('[a-zA-Z0-9’!"#$%&\'()*+,-./；:;<=>?@，：。?★、…【】《》？“”‘’！[\\]^_`{|}~\s一二三四五六七八九十]+', "",name)  # 事务所名称
            name = re.sub('人', '',name)
            name = re.sub('（）', '',name)
            shiwusuo_name.append(name)
            
            #获得事务所人数
            shiwusuo_number.append(re.search(r'[0-9]+',sws[4]).group(0))
            
            #获得事务所分类
            if '分所' in temp_sws[0] or '分公司' in temp_sws[0]:
                shiwusuo_fensuo.append('分所')
            else:
                shiwusuo_fensuo.append('总所')
            
            #获得负责人姓名（如有）
            if '主任会计师' in sws[1] :
                index = sws[1].index('主任会计师') 
                shiwusuo_manager.append(sws[1][index+7:index+10].replace(' ',''))
            elif '合伙人' in sws[1] :
                index = sws[1].index('合伙人')
                shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
            elif '负责人' in sws[1]:
                index = sws[1].index('负责人')
                shiwusuo_manager.append(sws[1][index+5:index+8].replace(' ',''))
            else:
                shiwusuo_manager.append('缺失')
            
            shiwusuo_single.append(shiwusuo_name)
            shiwusuo_single.append(shiwusuo_number)
            shiwusuo_single.append(shiwusuo_fensuo)
            shiwusuo_single.append(shiwusuo_manager)
            if '执业证书' not in sws[1]: #检查非姓名信息是否存于table中
                shiwusuo_single.append('缺失')
            final_shiwusuo.append(shiwusuo_single)
            
        else:
            print('出现异常情况')
    return final_shiwusuo

def findDropInformation(part1:list,table:list):
    temp_table1 = table
    drop_table = [] #Part1中缺失的信息，一个List代表一个Table缺失信息
    drop_number = 0 #记录需要删除的无效table数
    drop_index = []
    
    for i,information in enumerate(part1):
        if len(information) == 5 :
            drop_index.append(i)
    
    for i,information in enumerate(temp_table1):
        if '执业证书' in information[0][0]:
            drop_table.append(information)
            temp_table1[i] = 'Null'
            drop_number = drop_number + 1
        elif '联系电话' in information[0][0]:
            temp_table1[i] = 'Null'
            drop_number = drop_number + 1
        
      
    for i in range(drop_number):
        temp_table1.remove('Null')
    
    for i in range(len(drop_table)):
        if len(drop_table[i][0]) == 4:
            
            if '主任会计师' in drop_table[i][0][2] :
                index = drop_table[i][0][2].index('主任会计师') 
                part1[drop_index[i]][3][0] = drop_table[i][0][2][index+7:index+10].replace(' ','')
                part1[drop_index[i]].remove('缺失')
            elif '合伙人' in drop_table[i][0][2] :
                index = drop_table[i][0][2].index('合伙人') 
                part1[drop_index[i]][3][0] = drop_table[i][0][2][index+5:index+8].replace(' ','')
                part1[drop_index[i]].remove('缺失')
            elif '负责人' in drop_table[i][0][2]:
                index = drop_table[i][0][2].index('负责人') 
                part1[drop_index[i]][3][0] = drop_table[i][0][2][index+5:index+8].replace(' ','')
                part1[drop_index[i]].remove('缺失')
        
        elif len(drop_table[i][0]) == 3:
            
            if '主任会计师' in drop_table[i][0][1] :
                index = drop_table[i][0][1].index('主任会计师') 
                part1[drop_index[i]][3][0] = drop_table[i][0][1][index+7:index+10].replace(' ','')
                part1[drop_index[i]].remove('缺失')
            elif '合伙人' in drop_table[i][0][1] :
                index = drop_table[i][0][1].index('合伙人') 
                part1[drop_index[i]][3][0] = drop_table[i][0][1][index+5:index+8].replace(' ','')
                part1[drop_index[i]].remove('缺失')
            elif '负责人' in drop_table[i][0][1]:
                index = drop_table[i][0][1].index('负责人') 
                part1[drop_index[i]][3][0] = drop_table[i][0][1][index+5:index+8].replace(' ','')
                part1[drop_index[i]].remove('缺失')
                
        else :
            print('缺失信息中出现问题')
    
    return temp_table1,part1

def concatNameTable(nametables):
    final_nametables = []
    i = 0
    while i < len(nametables):
        if i != len(nametables)-1:
            if '姓名' in nametables[i][0][0].replace(' ','') and '姓名' in nametables[i+1][0][0].replace(' ',''):
                final_nametables.append(nametables[i])
                
            elif '姓名' in nametables[i][0][0].replace(' ','') and '姓名' not in nametables[i+1][0][0].replace(' ',''):
                j = 0
                temp_table = nametables[i]
                
                for table in nametables[i+1:len(nametables)]:
                    if '姓名' not in table[0][0].replace(' ',''):
                        j = j + 1
                        
                        for row in table:
                            temp_table.append(row)
        
                    else:
                        break
                        
                final_nametables.append(temp_table)
                i = i + j + 1
                continue
            else:
                print('拼接出现问题' + str(i))
        
        else:
            final_nametables.append(nametables[i])
        i = i + 1
            
    for row,i in enumerate(final_nametables):
        number = 0
        for j in range(0,len(i[len(i)-1]),2):
            if i[len(i)-2][j] == i[len(i)-1][j]:
                final_nametables[row][len(i)-1][j] = '重复'
                final_nametables[row][len(i)-1][j+1] = '重复'
                number = number + 2
        for n in range(number):
            final_nametables[row][len(i)-1].remove('重复')
    
    return final_nametables

def textNameNumber(final_name,final_part):
    
    for i,name_table in enumerate(final_name):
        name_number = 0
        for j,name_row in enumerate(name_table):
            if j != 0:
                name_number = name_number + (len(name_row)/2)
            else:
                continue
        if name_number != int(final_part[i][1][0]):
            print(name_number)
            print(i)
            break
        else:
            continue

def date(stamp):
    """输出时间格式
    """
    delta = pd.Timedelta(str(stamp)+'D')
    real_time = pd.to_datetime('1899-12-30') + delta
    return real_time

def word_extraction_Hunan_table(province:str,year:str):
    
    path = r'E:\Spyderworkplace\excel\湖南\2019_pdf'
    path_output = r'E:\Spyderworkplace\excel\湖南\2019_pdf'
    filelist = os.listdir(path)
    
    """ 生成excel的表头
    """
    i_out = 1
    workbook = xlwt.Workbook(encoding='utf-8')
    worksheet = workbook.add_sheet('sheet1', cell_overwrite_ok=True)
    worksheet.write(0, 0, label='年度')
    worksheet.write(0, 1, label='通知发布日期')
    worksheet.write(0, 2, label='来源公告名称')
    worksheet.write(0, 3, label='事务所所在省份')
    worksheet.write(0, 4, label='事务所全称')
    worksheet.write(0, 5, label='总所or分所')
    worksheet.write(0, 6, label='总/分所所长（如有）')
    worksheet.write(0, 7, label='注册会计师姓名')
    worksheet.write(0, 8, label='注册会计师ID')
    worksheet.write(0, 9, label='备注')
    
    for file in filelist:
        if '.docx' in file and year in file and province in file:  # 搜索docx
            word_path = os.path.join(path, file)
            tables = tableToList(word_path)
    
            paragraph = paragraphToList(word_path)
            title,index = getTitle(paragraph)
            shiwusuo = getShiwusuo(paragraph, index)
            part = getPart1(shiwusuo)
            nametables,final_part = findDropInformation(part,tables)   
            final_name = concatNameTable(nametables)
            
            wb_read = xlrd.open_workbook(os.path.join(path,'源地址.xls'))  # 从源地址excel提取'通知发布日期' '来源公告名称'
            
            for i in range(wb_read.sheet_by_index(0).nrows - 1):
                if year in wb_read.sheet_by_index(0).cell_value(i, 0) and province in wb_read.sheet_by_index(
                        0).cell_value(i, 0):
                    d = wb_read.sheet_by_index(0).cell(i, 1).value
                    date_pb = str(date(d)).split(" ")[0]
                    name_pb = wb_read.sheet_by_index(0).cell_value(i, 0)


            for st,x in enumerate(final_part):
                
                name_row,id_row = getTempName(final_name[st])

                for y in range(len(name_row)):
                    worksheet.write(i_out, 0, year)
                    worksheet.write(i_out, 1, date_pb)
                    worksheet.write(i_out, 2, name_pb)
                    worksheet.write(i_out, 3, province)
                    worksheet.write(i_out, 4, x[0][0])  # 事务所
                    if '总所' in x[2][0]:
                        worksheet.write(i_out, 5, '总所')
                    elif '分所' in x[2][0] or '分公司' in x[2][0]:
                        worksheet.write(i_out, 5, '分所')
                    else:
                        worksheet.write(i_out, 5, '总所')
                    worksheet.write(i_out,6,x[3][0])
                    worksheet.write(i_out, 7, name_row[y])
                    worksheet.write(i_out,8,id_row[y])
                    i_out = i_out + 1

                    
    workbook.save(os.path.join(path_output,province+year+'.xls'))  

def getTempName(name_table:list): #用于提取每个NameTable中会计师的姓名和ID，一个事务所会计师姓名转化成一个List
    
    single_shiwusuo_name = []
    single_shiwusuo_id = []
    
    for i in range(1,len(name_table)):
        for j in range(0,len(name_table[i])-1,2):
            single_shiwusuo_name.append(name_table[i][j])
            
    for i in range(1,len(name_table)):
        for j in range(1,len(name_table[i]),2):
            single_shiwusuo_id.append(name_table[i][j])
    
    return single_shiwusuo_name,single_shiwusuo_id
    
if __name__ == '__main__':
    
    '''
    path = '附件：湖南省注册会计师协会关于我省注册会计师2019年度任职资格检查合格人员名单（第一批）.docx'
    tables = tableToList(path)
    
    paragraph = paragraphToList(path)
    title,index = getTitle(paragraph)
    shiwusuo = getShiwusuo(paragraph, index)
    part1 = getPart1(shiwusuo)
    
    nametables,final_part1 = findDropInformation(part1,tables)   
    final_name = concatNameTable(nametables)
    textNameNumber(final_name, final_part1)
    '''
    word_extraction_Hunan_table('湖南','2019')
